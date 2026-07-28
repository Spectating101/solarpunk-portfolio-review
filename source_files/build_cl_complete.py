from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom'):
        if edge in kwargs:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(kwargs[edge]))
            el.set(qn('w:color'), '000000')
            tcBorders.append(el)
    tcPr.append(tcBorders)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(11.5)
style.paragraph_format.space_after = Pt(8); style.paragraph_format.line_spacing = 1.4
for s in doc.sections:
    s.left_margin = Inches(1); s.right_margin = Inches(1); s.top_margin = Inches(1); s.bottom_margin = Inches(1)

def add_chapter(text):
    h = doc.add_paragraph(); run = h.add_run(text)
    run.font.name='Times New Roman'; run.font.size=Pt(16); run.font.bold=True
    h.paragraph_format.space_before=Pt(6); h.paragraph_format.space_after=Pt(14)
    return h

def add_heading(text, size=12.5):
    h = doc.add_paragraph(); run = h.add_run(text)
    run.font.name='Times New Roman'; run.font.size=Pt(size); run.font.bold=True; run.font.color.rgb=RGBColor(0,0,0)
    h.paragraph_format.space_before=Pt(14); h.paragraph_format.space_after=Pt(6)
    return h

def add_para(text, bold=False, italic=False, size=11.5, align=None, space_after=8):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic
    p.paragraph_format.space_after=Pt(space_after)
    if align: p.alignment = align
    return p

def add_quote(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.italic=True
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text=''; p=hdr[i].paragraphs[0]; r=p.add_run(h)
        r.font.bold=True; r.font.size=Pt(9.5); r.font.name='Times New Roman'
        set_cell_border(hdr[i], top=8, bottom=4)
    for row in rows:
        cells = t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; p=cells[i].paragraphs[0]; r=p.add_run(str(val))
            r.font.size=Pt(9.5); r.font.name='Times New Roman'
    for c in t.rows[-1].cells: set_cell_border(c, bottom=8)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return t

def add_figure(path, caption, width=5.8):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.font.size=Pt(9.5); r.font.italic=True; r.font.name='Times New Roman'
    cap.paragraph_format.space_after=Pt(14)

FIG = "/tmp/cl_figures"

# ===================== TITLE PAGE =====================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("The Constrained Ledger:\nWhen Does Energy Actually Discipline Digital Money?")
r.font.size=Pt(18); r.font.bold=True; r.font.name='Times New Roman'
tp.paragraph_format.space_after=Pt(20)
for line,sz,bold in [("Christopher Ongko (王新福)",13,True),
    ("Department of Finance, College of Management, Yuan Ze University, Taiwan",11,False),
    ("s1133958@mail.yzu.edu.tw · ORCID: 0009-0007-9339-9098",11,False),
    ("Working draft · July 2026",11,False)]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=p.add_run(line); rr.font.size=Pt(sz); rr.font.bold=bold; rr.font.name='Times New Roman'
    p.paragraph_format.space_after=Pt(4)
doc.add_page_break()

# ===================== ABSTRACT =====================
add_heading("Abstract", size=14)
for para in [
"Bitcoin was engineered on a specific claim: that a currency's value could be disciplined by real, verifiable cost rather than institutional trust. This thesis tests whether that claim was ever actually delivered, using China's June 2021 mining ban as the clearest available natural experiment. A widely cited ratio of market value to cumulative mining cost (Hayes, 2019) is tested against four falsifiable criteria stated in advance: it must survive substitution with unrelated cumulative variables, survive a change in starting assumptions, hold in changes as well as levels, and be confirmed by a correctly specified structural-break test. The ratio fails all four. It is nearly indistinguishable from a ratio built on cumulative electricity use or cumulative elapsed time (correlation of .999989), depends materially on an arbitrary pre-sample accounting choice, disappears in differenced form, and receives no support from the preferred joint robust break test.",
"This negative result motivates the thesis's constructive contribution. Five conditions are specified as jointly necessary for an energy-linked financial claim to be genuinely constrained: reliable evidence, rule-bound issuance, explicit pricing of uncertainty, defined settlement, and limited governance. A transparent pricing framework, calibrated to five years of NASA satellite irradiance data, converts renewable-energy uncertainty into inspectable option prices, margin requirements, and data-tolerance limits. A working rule system implementing all five conditions is then tested against the same four-part standard that eliminated the passive shortcut: it is evaluated across differentiated cases, under a deliberate counterfactual, under changing settlement conditions, and through execution on a public blockchain test network. The system satisfies all four criteria, producing correct, traceable outcomes throughout.",
"The thesis does not claim to have built a functioning currency. Circulation, adoption, and liquidity are not tested or claimed anywhere in this work. What is demonstrated is the piece that neither fiat money nor Bitcoin has actually delivered: a financial constraint tied to real evidence that can be inspected and shown to hold under test, rather than taken on faith.",
]:
    add_para(para)
add_para("Keywords: energy-linked finance, digital money, monetary credibility, Bitcoin energy cost, renewable-energy risk, smart contracts, reproducible decision systems", italic=True, size=10)
add_para("JEL Codes: E42, G13, Q42, Q47", italic=True, size=10)
doc.add_page_break()

# ===================== TABLE OF CONTENTS (simple) =====================
add_heading("Table of Contents", size=14)
toc_entries = [
    "Chapter 1 — Introduction",
    "Chapter 2 — Why an Energy Reference Is Not Automatically a Constraint",
    "Chapter 3 — Testing Bitcoin as the Clearest Case: Method",
    "Chapter 4 — Results and Interpretation",
    "Chapter 5 — Pricing Renewable-Energy Uncertainty",
    "Chapter 6 — The Built, Tested System",
    "Chapter 7 — Discussion and Conclusion",
    "Appendix A — Supplementary Empirical Checks",
    "Appendix B — Implementation Evidence and Source Lineage",
    "References",
]
for e in toc_entries:
    p = doc.add_paragraph(); r = p.add_run(e)
    r.font.name='Times New Roman'; r.font.size=Pt(11.5)
    p.paragraph_format.space_after=Pt(6)
doc.add_page_break()

# ===================== CHAPTER 1 =====================
add_chapter("Chapter 1 — Introduction")

add_heading("1.1 Research Motivation")
add_para("Bitcoin was engineered, not only launched, on a specific claim: that a currency's value could be disciplined by real, verifiable cost rather than institutional trust. In June 2021, that claim faced its clearest available test. China's State Council banned cryptocurrency mining nationwide, forcing roughly seventy percent of Bitcoin's global computing power to relocate within months. If cumulative mining cost genuinely anchors Bitcoin's price the way its engineering was intended to guarantee, this shock should have produced a detectable change in that relationship.")
add_para("A specific ratio — market value divided by cumulative mining expenditure — has circulated across crypto analysis, and is defended in published research (Hayes, 2019), as evidence that this claim was correct: that mining cost functions as a price floor. Tested directly, using daily data from 2019 through 2025, the relationship does not survive scrutiny. A formal test for whether it changed after the ban finds no reliable break. Replacing cumulative electricity cost with an unrelated variable — cumulative elapsed time — produces a nearly identical result, indicating the ratio was never capturing information specific to electricity price at all. This finding is the starting point for the analysis that follows.")
add_para("Bitcoin's own founding document explicitly stated its ambition: its title called it “a peer-to-peer electronic cash system” — a currency, declared as a goal, years before any real adoption existed to support that claim. That same ambition motivates the analysis that follows: not a narrower financial instrument for its own sake, but the objective Bitcoin itself declared and, as this thesis demonstrates, did not deliver — a currency system whose value is genuinely disciplined by something real and checkable, rather than by an institution's word or an unverified mechanism. The ambition is not achieved in full here. No claim is made about circulation, adoption, or liquidity — Bitcoin established its own currency status empirically, over fifteen years of real use, rather than by assertion in a paper. What is demonstrated is the specific requirement needed to make that ambition achievable at all, tested directly for the first time. Chapter 7 addresses how far that gets, and what would still be required before this could honestly be called a currency system rather than evidence that one might be possible.")

add_heading("1.2 Problem Statement")
add_para("An energy reference does not, by itself, constrain anything. A fixed token supply does not reveal whether the underlying evidence is reliable, whether the permitted quantity is defensible, or whether settlement obligations can be met — an energy label can coexist with weak data, discretionary issuance, and no guarantee of what happens when a claim comes due. Bitcoin's cumulative mining cost, tested directly in this thesis, is the clearest available case of this gap: real expenditure, with no delivered constraint. The problem this thesis addresses is what would actually be required to close that gap.")

add_heading("1.3 Research Question")
add_para("This thesis addresses one question directly:")
add_quote("Under what conditions can verified energy evidence impose a credible, checkable constraint on a digital financial claim?")
add_para("Three supporting questions test that proposition: (1) Does Bitcoin's mining-cost ratio identify a value effect specific to energy cost, or to persistence and construction artifacts alone? (2) How can renewable-energy uncertainty be translated into explicit, inspectable pricing and risk quantities? (3) Can evidence quality, admission, quantity limits, and settlement be represented as separate, reproducible decision stages, and do they behave correctly under test?")

add_heading("1.4 Research Design")
add_para("The thesis proceeds through three connected stages. The first tests the strongest available shortcut — that Bitcoin's cumulative mining cost alone anchors its value — against a real natural experiment, using negative controls and robustness checks designed to distinguish a genuine electricity-price effect from statistical artifact. The second translates renewable-energy uncertainty into concrete financial quantities using a transparent, publicly calibrated pricing model. The third specifies and implements the conditions under which those quantities could actually govern a real claim, testing the resulting system against controlled cases with differentiated, traceable outcomes.")

add_heading("1.5 Contributions")
add_para("This thesis makes three contributions.")
add_para("First, it identifies the limits of Bitcoin's cumulative mining cost as a value anchor: the relationship a widely cited ratio (Hayes, 2019) claims to show is not specific to electricity price, does not survive a proper structural-break test, and disappears under differencing — a negative result reached through several independent checks rather than one specification.")
add_para("Second, it builds a transparent pricing framework for renewable-energy uncertainty, calibrated to five years of satellite irradiance data, converting a physically variable resource into inspectable option prices, margin requirements, and data-tolerance limits.")
add_para("Third, it specifies five jointly necessary conditions for an energy-linked financial claim to be credible — reliable evidence, rule-bound issuance, explicit pricing of uncertainty, defined settlement, and limited governance — and implements and tests a system built on them, including execution on a public blockchain and controlled test cases with traceable, differentiated outcomes.")

add_heading("1.6 Scope of the Thesis")
add_para("The claims here are deliberately bounded. This thesis does not argue that energy should replace fiat money, does not treat Bitcoin as energy-backed, and does not present the implemented system as legally or commercially ready. Three boundaries define the evidence: (1) the Bitcoin analysis identifies the limits of one mining-cost ratio, not a general law of digital-asset valuation; (2) the renewable-energy model is a cold-start framework for exposing pricing and risk assumptions, not a traded price or legal settlement mechanism; (3) the implementation is proof-of-concept research on a public testnet, not evidence of operator validation or deployment readiness.")
doc.add_page_break()

# ===================== CHAPTER 2 =====================
add_chapter("Chapter 2 — Why an Energy Reference Is Not Automatically a Constraint")
add_para("Chapter 1 argued that fiat and Bitcoin each carry a specific, known limitation, and that neither delivers money genuinely disciplined by something checkable. This chapter grounds that claim in the relevant literature, and specifies what would be required to achieve what neither model has done.")

add_heading("2.1 Credible Constraint as the Organizing Question")
add_para("The existing literature does not lack pieces of this puzzle. It offers explanations for scarcity, credibility, production risk, and automated enforcement. What it lacks is an account of how those pieces interact before a physical reference can actually limit a digital financial claim. Monetary economics explains why commitment and limits on discretion matter (Barro & Gordon, 1983; Kydland & Prescott, 1977). Renewable-energy finance describes productive but genuinely variable output (International Energy Agency, 2023; Lazard, 2025). Pricing theory forces uncertainty into stated, checkable assumptions, and smart contracts can enforce whatever rules are declared (Black & Scholes, 1973; Cong & He, 2019; Cox, Ross, & Rubinstein, 1979). Each literature addresses one part of the design problem, and none addresses all of it. The organizing question is not which literature supplies the correct metaphor for “energy money,” but which combination of commitment, evidence, pricing, settlement, and governance actually turns an energy reference into a real constraint.")

add_heading("2.2 Monetary Credibility, Rules, and Discretion")
add_para("Kydland and Prescott (1977) demonstrate that a rule announced today can be abandoned once people have already acted on it, so credibility depends not only on the current rule but on how costly and visible it would be to change it later. Barro and Gordon (1983) identify the related expectations problem: promises are discounted whenever people expect future discretion. In digital finance, a claim of limited supply or energy backing remains weak for the same reason if administrators can mint outside the rule, replace the evidence source, or avoid redemption when inconvenient. The design principle this yields is not a specific asset choice — it is that the limits must bind the issuer as much as the user.")

add_heading("2.3 Gold and Fiat")
add_para("Gold and fiat each ground Chapter 1's comparison in monetary history. Gold-backed money combined physical scarcity with a working redemption promise; its most visible failure, at Bretton Woods, was not gold disappearing — it was the arithmetic connecting outstanding claims to settlement capacity breaking down (Bordo, 1993; Eichengreen, 1992). Fiat money replaced that physical promise with institutional authority and a track record, trading flexibility for dependence on political will and repeated performance (Federal Reserve Bank of St. Louis, 2010). Energy differs from both in ways that matter throughout the remainder of this thesis: it is time- and location-dependent, often impossible to store economically, and inseparable from the infrastructure needed to deliver it. An energy-linked system cannot import gold's convertibility or fiat's institutional depth by analogy — it must define its own evidence, pricing, settlement, and governance.")

add_heading("2.4 Bitcoin, Protocol Scarcity, and Proof-of-Work")
add_para("Bitcoin's issuance schedule is fixed by protocol, and proof-of-work makes rewriting transaction history costly (Nakamoto, 2008) — it is neither redeemable like gold nor administered like fiat. Proof-of-work provides digital scarcity with a real expenditure base: miners consume electricity and hardware to compete for rewards, making Bitcoin genuinely costly to produce and secure. Costly production is not redeemable backing, however — holders receive no right to the electricity actually consumed (de Vries, 2018). This makes Bitcoin a useful boundary case: any distinctive mining-expenditure effect on price must be separated from ordinary crypto return dynamics, price persistence, and how the ratio itself is constructed (Granger & Newbold, 1974; Kristoufek, 2015; Kronmal, 1993; Liu & Tsyvinski, 2021) — exactly what Chapter 3 tests.")

add_heading("2.5 Bitcoin Energy Valuation and Cryptocurrency Returns")
add_para("Hayes (2019) reports evidence consistent with a marginal production-cost view of Bitcoin's price — the result that motivates Chapter 3's test directly, though it does not by itself establish a guaranteed value floor. Mining cost changes with electricity prices, hardware efficiency, and network difficulty, while market value responds to demand, liquidity, and coordination, so a production-cost association must be identified rather than inferred from the fact that mining is expensive. The available data has real limits: the Cambridge Bitcoin Electricity Consumption Index is model-based, not direct metering (Cambridge Centre for Alternative Finance, n.d.-a, n.d.-c), and any cumulative-cost series inherits that uncertainty — which is why Chapter 3 treats CEIR as a falsifiable test rather than a measure of intrinsic value.")

add_heading("2.6 Renewable-Energy Finance, Data, and Risk")
add_para("Solar and wind assets generate electricity with direct economic use, but the financial value of that output depends on weather, location, timing, grid access, and settlement rules (Joskow, 2011; Ueckerdt et al., 2013). Cost competitiveness does not eliminate those contingencies (Lazard, 2025). Public data illustrates a real evidence hierarchy: NASA POWER and PVWatts are useful for screening and scenario analysis, but neither can prove that a specific site generated, exported, or settled a given quantity of electricity (Dobos, 2014; Sengupta et al., 2024) — resource potential, expected output, observed generation, and settled delivery are genuinely different things, and Chapter 5 treats them accordingly.")

add_heading("2.7 Pricing Energy-Linked Claims")
add_para("Black and Scholes (1973) and Cox, Ross, and Rubinstein (1979) are used here not because electricity behaves like an ordinary stock, but because they require the analyst to declare the underlying proxy, volatility, horizon, and payoff explicitly. Standard option assumptions fit electricity imperfectly, since power prices can spike, mean-revert, and turn negative (Lucia & Schwartz, 2002), so option-style reasoning is used in Chapter 5 as a transparent benchmark rather than a complete market model.")

add_heading("2.8 Smart Contracts, Oracles, and Governance")
add_para("Smart contracts can enforce issuance and settlement rules, but execution quality is not economic quality — automation can apply a flawed rule or accept inaccurate data as faithfully as sound ones (Cong & He, 2019). The part that matters directly for what follows is the oracle problem: a blockchain cannot itself observe whether electricity was generated or delivered, so external data feeds must supply that fact, introducing their own reliability and manipulation risk (Chainlink, 2025; Eskandari et al., 2021; Zhang et al., 2016) — a limitation Chapter 7 addresses directly. Governance introduces a parallel risk: an administrator can defeat every other safeguard without falsifying the data, simply by minting outside accepted evidence (Kiayias & Lazos, 2022). This framework is not a stablecoin design — it addresses an earlier question than reserve policy or peg mechanics: whether verified evidence can restrict claim creation at all (Bank for International Settlements, 2023).")

add_heading("2.9 Research Gap and Five Constraints")
add_para("Credibility, across this literature, emerges from a chain rather than from any single asset or technology. The unresolved gap lies in the connection between these pieces: Bitcoin research studies mining expenditure without connecting it to a direct energy claim; renewable-energy finance studies production without specifying issuance rules; pricing research studies uncertain payoffs without specifying how the result should limit claim creation; smart-contract research studies enforcement without determining whether the underlying claim is economically defensible. The five constraints introduced in Chapter 1 — reliable evidence, rule-bound issuance, explicit pricing of uncertainty, defined settlement, and limited governance — are this thesis's answer to that gap, and the remaining chapters test, price, and implement them directly.")
doc.add_page_break()

# ===================== CHAPTER 3 =====================
add_chapter("Chapter 3 — Testing Bitcoin as the Clearest Case: Method")

add_heading("3.1 The Testable Claim")
add_para("This chapter tests the shortcut that motivates the thesis: whether real mining expenditure already anchors Bitcoin's value on its own. Bitcoin is the relevant case because proof-of-work converts electricity and computation into one measurable, cumulative cost.")
add_para("For that shortcut to count as genuine evidence of an energy-price constraint, rather than a suggestive number, it must clear four specific tests, stated here before any result is reported:")
add_para("1. It must survive substitution. Replacing cumulative mining cost with an unrelated cumulative quantity — electricity used rather than its price, or simply elapsed time — should weaken or eliminate the result. If the same pattern survives the substitution, the original result was never specific to energy cost.")
add_para("2. It must survive a change in starting assumptions. If removing an arbitrary pre-sample accounting choice materially changes the result, the result rested on that choice rather than the underlying relationship.")
add_para("3. It must hold in changes, not only in levels. A genuine causal relationship should appear in short-run movements as well as cumulative totals. A result that appears only in slow-moving levels is the signature of two persistent series drifting together, not evidence of one driving the other.")
add_para("4. It must be confirmed by the correctly specified test, not only the simpler one. Where a classical test and a properly robust joint test disagree, the more carefully specified test governs the conclusion.")
add_para("This is not an arbitrary checklist. It is the standard any claim of genuine constraint must meet — not only a ratio, but any system proposing that evidence disciplines a financial claim. Chapters 5 and 6 hold the system this thesis builds to an analogous standard: whether it behaves correctly under substitution across different cases, under a deliberate change to a single input, and under different real conditions at settlement. What follows is the first application of that standard, to the most direct candidate available.")

add_heading("3.2 Why Bitcoin Is the Relevant Test Case")
add_para("Bitcoin separates expenditure from claim cleanly: mining consumes real electricity, but holders have no redeemable right to it, which makes Bitcoin a useful test of whether costly production leaves a distinctive valuation trace regardless. Three features make it informative. First, mining produces a cumulative expenditure series directly comparable to market capitalization — negative controls are necessary because both series are highly persistent, and persistence alone can manufacture an apparent relationship. Second, China's 2021 mining restrictions materially changed where hashrate was located (Cambridge Centre for Alternative Finance, n.d.-b) — a plausible candidate break, not proof that the cost mechanism itself changed. Third, the regime interpretation is genuinely testable: separate regressions and the classical Chow statistic describe the raw difference, while a joint robust test evaluates whether the coefficients changed together with statistical precision (Chow, 1960). The working hypothesis — that concentrated mining might create a common cost reference, and later dispersion could weaken it — is treated as something to be tested, not assumed.")

add_heading("3.3 Measuring the Relationship")
add_para("The test is organized around the Cumulative Energy Investment Ratio:")
add_quote("CEIRₜ = Market Capitalisationₜ / Cumulative Energy Costₜ")
add_para("A high value indicates Bitcoin's market value is large relative to estimated cumulative mining cost; the candidate sign is negative. Because the denominator accumulates slowly and persistently, this interpretation is vulnerable to spurious regression (Granger & Newbold, 1974; Kronmal, 1993) — market capitalization carries the live market-price signal in the numerator, while the denominator simply accumulates. An energy-specific result should outperform ratios built on cumulative electricity use or cumulative elapsed days, and remain stable when the starting cost assumption or break specification changes.")
add_para("The panel runs from January 2019 through May 2025. Bitcoin price and market capitalization are derived from a local daily price export and an approximate supply curve rather than an observed market-cap series. Mining electricity use comes from the Cambridge Bitcoin Electricity Consumption Index, a model-based estimate rather than direct miner metering (Cambridge Centre for Alternative Finance, n.d.-a). The intended geography-weighted electricity-price series did not merge as designed, leaving a near-constant price path in its place — addressed directly in Appendix B. The candidate split date, June 20, 2021, is a researcher-defined proxy for the broader China mining crackdown.")

add_heading("3.4 The Empirical Design")
add_para("The design separates three questions: prediction, regime change, and mechanism. The first test asks whether CEIR predicts Bitcoin's next 30-day return, using standard errors designed for overlapping return windows (HAC(30); Hodrick, 1992; Newey & West, 1987). The second asks whether the relationship changes around the candidate split, using a joint robust test as the preferred measure alongside the classical Chow statistic for comparison (Chow, 1960). The third asks whether any apparent effect is specific to energy cost — removing the starting cost assumption, substituting cumulative electricity use and elapsed days for cumulative cost, estimating the relationship in changes rather than levels, and checking for a genuine long-run relationship rather than a coincidental one (Dickey & Fuller, 1979; Engle & Granger, 1987).")

add_heading("3.5 Data and Sample")
add_para("The regression sample contains 2,280 daily observations once forward returns and controls are complete, drawn from a full panel of 2,340 days — 901 observations before the candidate split, 1,439 after.")
add_table(
    ["Series", "Source", "Frequency", "Role"],
    [["Bitcoin price and derived market capitalization", "Local daily price export; approximate BTC supply curve", "Daily", "Return outcome; CEIR numerator"],
     ["Mining electricity", "Cambridge CBECI best-guess annualized TWh", "Daily model estimate", "Energy-cost base"],
     ["Legacy electricity-price input", "Intended geography-weighted merge; failed, near-constant instead", "Daily", "Legacy cost denominator"],
     ["Mining geography", "Cambridge Mining Map, monthly shares", "Monthly", "Context for the candidate split"],
     ["Fear & Greed Index", "Alternative.me composite", "Daily", "Sentiment control"]]
)
add_para("Table 3.1. Data sources.", italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table(
    ["Variable", "Mean (full)", "Mean (pre-split)", "Mean (post-split)"],
    [["Bitcoin price (USD)", "34,881", "16,666", "46,161"],
     ["CEIR", "29.65", "30.40", "29.19"],
     ["log(CEIR)", "3.28", "3.29", "3.28"],
     ["30-day forward return", "6.4%", "10.7%", "3.8%"],
     ["30-day return volatility", "3.2%", "3.7%", "2.8%"]]
)
add_para("Table 3.2. Descriptive statistics, regression sample.", italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Bitcoin's mean price roughly tripled after the split, while mean CEIR remained in a broadly similar range — the two series do not move in lockstep, which is part of why formal inference, rather than the raw picture, must carry the argument. A simple correlation between log(CEIR) and forward returns is negative (-0.19) even before any of the formal tests in Chapter 4 are run — consistent with the attractive pre-split story, and the number Chapter 4's four tests exist to interrogate rather than accept at face value.")

add_heading("3.6 Anticipated Interpretation of Results")
add_para("Before Chapter 4 reports results, it is worth stating what each outcome would mean. If CEIR clears all four tests, the reading is that Bitcoin's price is genuinely disciplined by what it cost to mine, and the China ban would represent a real disruption to that discipline. If CEIR fails the substitution test specifically, the reading is that the ratio reflects a general valuation pattern rather than an electricity-specific one. Failing the starting-assumption test indicates the result depends on bookkeeping rather than economics. Failing the changes-versus-levels test points to persistence rather than causation. Failing the break test means the China ban — the one real natural experiment available — found nothing. The next chapter reports which of these outcomes actually occurred.")
doc.add_page_break()

# ===================== CHAPTER 4 =====================
add_chapter("Chapter 4 — Results and Interpretation")

add_heading("4.1 CEIR and Its Failure")
add_para("The level regression reproduces the attractive result on its face: before the candidate split, a higher CEIR predicts weaker 30-day returns, with a coefficient of -0.262 (HAC p < .001) — a strong-looking relationship by conventional standards. Table 4.1 reports what happens to that result under each of Chapter 3's four tests.")
add_table(
    ["Test", "What was checked", "Result", "Reading"],
    [["1. Substitution", "Swap cumulative cost for cumulative electricity use or cumulative days", "Correlation of .999989 with the electricity-use version; similarly strong with the day-count version", "Fails — not specific to price"],
     ["2. Starting assumptions", "Remove ~$3.30B of pre-sample 2018 cost", "Coefficient weakens from -0.262 to -0.104; p-value rises from <.001 to .099", "Fails — depends on an arbitrary choice"],
     ["3. Changes vs. levels", "Re-estimate using period-to-period changes", "No reliable predictive relationship in changes", "Fails — consistent with persistence"],
     ["4. Correct break test", "Classical Chow statistic vs. joint robust test", "Chow suggests a break; joint robust test does not (p = .133)", "Fails to confirm a break"]]
)
add_para("Table 4.1. Results against the four-part standard.", italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table(
    ["Test", "N", "Estimate", "Robust p", "Reading"],
    [["Pre-split CEIR", "872", "β = -0.2623", ".00052", "Real association"],
     ["Post-split CEIR", "1,408", "β = -0.0708", ".133", "Weak"],
     ["Joint robust break test", "2,280", "—", ".133", "No reliable break"]]
)
add_para("Table 4.2. Primary regression results.", italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Two further checks — a differenced specification and a trading-rule evaluation — confirm the same conclusion and are reported in full in Appendix A. CEIR fails all four tests. Mining expenditure is real; this ratio does not isolate its financial effect.")

add_heading("4.2 Interpretation")
add_para("This negative result rules out an unsupported shortcut, and it is worth stating precisely why. Even had CEIR cleared every one of Chapter 3's four tests, it would still have established only one thing: that market value moves inversely with cumulative mining cost. A single number of that kind cannot specify an evidence standard, an issuance rule, a settlement obligation, or a limit on governance. Chapter 2 already established that a real constraint requires all five conditions — reliable evidence, rule-bound issuance, explicit pricing of uncertainty, defined settlement, and limited governance — held jointly, not inferred from a ratio's sign. Even a passing CEIR would have been, at most, indirect evidence that discipline was present. It would not have been the discipline itself.")
add_para("That distinction is what CEIR's failure clarifies. It does not only rule out one specific ratio — it rules out the entire category of approach the ratio represents: inferring a constraint from an observed statistical pattern rather than building one directly out of evidence and rules. What remains is the possibility Chapter 1 opened with: that a constraint could be built deliberately, and tested, rather than inferred after the fact.")

add_heading("4.3 Implications for the Remainder of the Thesis")
add_para("Satisfying the five conditions requires two components that do not yet exist at this point in the thesis, each serving a distinct purpose. The first is a method for pricing genuine uncertainty in renewable-energy output — what constraint three, explicit pricing of uncertainty, actually requires, and something nothing built so far in this thesis provides: Bitcoin's cumulative cost is a single historical number, not a forward-looking price for something that varies by weather and season. The second is an implemented rule system that checks evidence, applies issuance limits, and handles settlement as separate, traceable decisions, addressing constraints one, two, four, and five together.")
add_para("Both components are held to the same discipline that eliminated CEIR: whether the result survives substitution across different cases, survives a deliberate change to a single input, and holds under different real conditions rather than only the case it was built around. Chapter 5 builds the first component and demonstrates that it holds up under its own version of that test — two independent pricing methods agreeing with each other. Chapter 6 builds and tests the second, against real cases with distinct outcomes, including a deliberate counterfactual designed specifically to check whether its decisions track evidence and rule rather than something else.")
doc.add_page_break()

# ===================== CHAPTER 5 =====================
add_chapter("Chapter 5 — Pricing Renewable-Energy Uncertainty")

add_heading("5.1 The Pricing Problem")
add_para("Ruling out the CEIR shortcut does not resolve the constructive problem it was standing in for. Renewable energy is genuinely productive, but its financial value depends on weather, location, and time of day — a solar farm does not produce the same output on consecutive days, and unlike Bitcoin's fixed, backward-looking cumulative cost, this is a forward-looking, ongoing uncertainty. Before constraint three from Chapter 2 can govern anything real, that variability must become an actual number. No liquid options market exists for this kind of contract against which a price could be checked, so this chapter builds a defensible number from first principles.")

add_heading("5.2 Calibration")
add_para("The method begins with five years of daily satellite-measured solar irradiance data (NASA POWER) for a representative Taiwan location. Raw daily irradiance is noisy — cloud cover and single-day weather variation would otherwise dominate the signal — so the data is smoothed with a four-day rolling average before daily log returns are computed, and the top and bottom one percent of those returns are trimmed to remove measurement outliers rather than genuine variability. This produces an annualized volatility estimate of approximately 189 percent — high, but this reflects direct physical output variability rather than financial market volatility, and the two should not be read as equivalent.")
add_para("That volatility feeds a standard option-pricing setup: Geometric Brownian Motion, priced two independent ways — a binomial tree (Cox, Ross, & Rubinstein, 1979) and a Monte Carlo simulation — chosen because both methods require every assumption (the underlying proxy, the horizon, the discount rate) to be stated explicitly, consistent with the role Chapter 2 assigns to pricing theory.")
add_table(
    ["Input", "Value"],
    [["Spot / strike price", "$0.0516/kWh"],
     ["Annualized volatility", "~189%"],
     ["Maturity", "1 year"],
     ["Binomial tree price (N=400 steps)", "$0.0356/kWh"],
     ["Monte Carlo price", "$0.0361/kWh"],
     ["Agreement between methods", "Within 2%"]]
)
add_para("Table 5.1. Taiwan base-case pricing.", italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("The two independent methods agreeing within two percent is this chapter's own consistency check — not proof the price is correct in a market sense, since no market yet exists to check it against, but evidence the calculation is stable rather than an artifact of one numerical approach. The same volatility estimate converts directly into a margin requirement of approximately $0.63/kWh, at a stress multiple of 1.5 times the estimated 99th-percentile daily loss, and a data-tolerance bound of approximately 21.7 percent — how much measurement error the system can absorb before it should stop trusting the underlying evidence.")

add_heading("5.3 What This Does and Does Not Establish")
add_para("This produces a defensible, inspectable set of numbers a policy can use to set a quantity limit or collateral requirement. It does not establish a true market price, since no liquid market yet exists to validate it against, and only Taiwan's volatility is calibrated from fully disclosed real data; applying this to another location would require independent calibration rather than a borrowed number. The pricing tool built here is an input to Chapter 6's system, not a market-validated instrument in its own right.")
doc.add_page_break()

# ===================== CHAPTER 6 =====================
add_chapter("Chapter 6 — The Built, Tested System")

add_heading("6.1 From Five Conditions to an Implemented System")
add_para("Chapter 2 specified five conditions a credible energy-linked claim requires. Chapter 5 built the pricing tool constraint three requires. This chapter implements both as a working rule system rather than leaving the conditions as description — separating four decisions that are easily blurred into a single opaque judgment: whether the evidence is trustworthy, whether policy allows the claim, how large it may be, and whether it can actually be settled. Each decision is made and recorded independently, so a later reviewer can trace which one produced a given outcome.")

add_heading("6.2 The Decision Pipeline")
add_figure(f"{FIG}/fig2_pipeline.png", "Figure 6.1. The four-gate decision pipeline, applied to three real test cases with three different outcomes.")
add_para("At the first gate, submitted evidence — a meter reading, a modeled production estimate, an attestation — is checked against a declared policy for what counts as acceptable, and a source hash is recorded so the same evidence cannot support a second claim. At the second gate, admission returns a plain PASS or BLOCK; nothing about size is calculated if this gate fails. At the third, quantity is capped by whichever applicable limit is lowest, expressed in the same units as the claim. At the fourth, settlement is evaluated later, against whatever capacity is actually available, rather than assumed automatically once admission has passed. A receipt preserves the inputs, the rule set used, whichever limit ultimately bound, the final outcome, and which version of the policy and data were in effect.")

add_heading("6.3 Testing the System Against Chapter 3's Standard")
add_para("This chapter applies the tests specified in Chapter 3 to the constructed system.")
add_table(
    ["Chapter 3's test", "Applied here as", "Result"],
    [["1. Survives substitution", "Tested across three different cases — Taoyuan, Austin, Phoenix", "Each produces a different, correctly traced outcome"],
     ["2. Survives a change in starting assumptions", "A deliberate counterfactual: only the declared assurance level changes, evidence file untouched", "Outcome changes from blocked to admitted, tracing to the rule that changed"],
     ["3. Holds under changing conditions", "Settlement tested against three capacity scenarios", "Correctly reports full settlement, partial settlement, and shortfall"],
     ["4. Confirmed by the correct test", "Real execution on a public blockchain test network, not a simulation", "Produces an independently checkable transaction"]]
)
add_para("Table 6.1. The four-part standard, applied to the constructed system.", italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Taoyuan is submitted with evidence that falls short of the required standard — the system blocks it outright, calculating no quantity, since there is nothing to size until the evidence clears the bar. In the counterfactual, the evidence file is left completely untouched and only the declared assurance level is raised to meet the policy's threshold; admission then passes, and the maximum allowed quantity — 126 units — traces directly to the provenance rule that changed, not to any change in the underlying evidence. Austin is admitted at 283.09811 units because the resource data itself is the lowest applicable ceiling; Phoenix is admitted at up to 320 units because a different, evidence-backed capacity constraint binds instead — two different reasons, not the same rule applied twice. Run against three settlement-capacity scenarios — 100 percent, 40 percent, and 0 percent — the system correctly reports SETTLED, PARTIAL, and SHORTFALL; at 40 percent capacity, it records 50.4 units covered and 75.6 units in shortfall, rather than rounding the gap away. One execution ran on Ethereum's Sepolia test network, minting SPK against 2,606.7 kWh of accepted surplus, producing a real, independently checkable transaction rather than a simulated one. At the reviewed revision, the deterministic core passed 60 of 60 tests and the front end passed 64 of 64.")
add_para("Where CEIR failed all four of the tests specified in Chapter 3, the system tested here satisfies the corresponding four criteria.")

add_heading("6.4 What This Establishes, and What It Does Not")
add_para("This establishes that evidence, issuance, pricing, settlement, and governance can be represented as separate, inspectable decisions rather than folded into a single opaque judgment — and that the separation changes outcomes, not only how they are explained. It does not establish that the underlying data sources are production-grade: satellite estimates and modeled production are not equivalent to metered output, and oracles that report real-world facts to a blockchain carry their own reliability and manipulation risk, as Chapter 2 notes directly. It does not establish legal enforceability, real market demand, or production-grade security. Consistent with the boundaries Chapter 1 set, this is proof-of-concept evidence under controlled conditions — Chapter 7 states directly what would still be needed to call it more than that.")
doc.add_page_break()

# ===================== CHAPTER 7 =====================
add_chapter("Chapter 7 — Discussion and Conclusion")

add_heading("7.1 Answer to the Research Question")
add_para("The central question was: under what conditions can verified energy evidence impose a credible, checkable constraint on a digital financial claim? The evidence developed across Chapters 3 through 6 supports a conditional answer. Passive expenditure alone does not constitute such a constraint — Chapter 4 shows that Bitcoin's cumulative mining cost fails every test that would distinguish a genuine energy-price effect from a persistent, coincidental valuation pattern. A constraint becomes credible only when five conditions are held jointly: reliable evidence, rule-bound issuance, explicit pricing of uncertainty, defined settlement, and limited governance. Chapters 5 and 6 demonstrate that these conditions can be implemented, and that the resulting system behaves correctly under the same discipline that eliminated the passive shortcut.")

add_heading("7.2 What the Results Establish")
add_para("Three findings support this answer. First, a negative result: the ratio most directly associated with the claim that mining cost anchors Bitcoin's value does not survive scrutiny once tested against substitution, changed assumptions, differencing, and a properly specified structural-break test. Second, a constructive tool: renewable-energy output can be converted into a defensible, internally consistent price using publicly available satellite data, without claiming to have discovered a true market price. Third, a tested system: the five conditions can be implemented as a working rule system, and that system produces correct, traceable outcomes across cases with different evidence, different binding constraints, and different settlement capacities, including one execution on a public blockchain producing a verifiable transaction.")

add_heading("7.3 Limitations")
add_para("Several limitations qualify these findings directly.")
add_para("The empirical evidence in Chapters 3 and 4 is limited to one proof-of-work asset, using a market capitalization series derived from a local price export and an approximate supply curve rather than a fully documented observed series. Mining electricity use is itself a model-based estimate, not direct metering, and the intended geography-weighted electricity-price series did not merge as designed. These limitations do not prevent reproduction of the original result; they limit how strongly that result, or its failure, can be generalized to a differently constructed cost measure.")
add_para("The pricing framework in Chapter 5 remains a cold-start scenario model. Real power markets exhibit jumps, negative prices, mean reversion, and seasonal patterns beyond what a Geometric Brownian Motion benchmark captures. Only the Taiwan case has a fully disclosed, preserved volatility calibration; other locations rely on illustrative inputs.")
add_para("The implementation in Chapter 6 is controlled proof-of-concept evidence. Execution on a public test network does not establish revenue-grade metering, legal enforceability, real market demand, or production-grade security. The satellite and model-based evidence used in the controlled cases is not equivalent to direct physical metering, and oracle risk — the possibility that external data feeds misreport or are manipulated — remains a real, unresolved boundary, as discussed in Chapter 2.")

add_heading("7.4 What Would Strengthen or Weaken This Work")
add_para("The architecture developed here is falsifiable, and it is worth stating plainly what would count against it. It would be weakened if real operator evidence could not be classified consistently under the stated policy, if the pricing outputs could not be converted into defensible quantity limits under real market data, if the same policy produced unpredictably different admission outcomes under equivalent evidence, if settlement rights remained ambiguous once tested against a real dispute, or if governance could waive the constraint whenever it became costly to honor. Conversely, the negative result on Bitcoin's mining cost would be weakened if a differently constructed, geography-accurate cost measure produced a result that survived the same four tests applied in Chapters 3 and 4.")

add_heading("7.5 Future Research")
add_para("A future extension would replace the model-based electricity-price series with a properly merged, geography-weighted series and repeat the identification tests in Chapters 3 and 4 under that correction. A pilot deployment, with real operator evidence and real settlement obligations, would test the implementation in Chapter 6 under conditions this proof-of-concept does not reach. Extending the pricing framework in Chapter 5 to incorporate jump-diffusion or mean-reverting dynamics would address the limitation acknowledged above.")

add_heading("7.6 Closing Statement")
add_para("China's 2021 mining ban was, in principle, the clearest available test of whether Bitcoin's cost of production disciplines its value. It does not. The relationship claimed by a widely cited production-cost ratio fails every test designed to distinguish a genuine energy-price effect from a coincidental one.")
add_para("What replaces that claim is not a theory. It is a system: five jointly necessary conditions, derived from how monetary discipline has actually functioned in gold-backed and fiat regimes, implemented as a working rule set, and tested against real cases with differentiated, traceable outcomes.")
add_para("Chapter 1 stated an ambition as large as Bitcoin's own founding claim: a currency genuinely disciplined by something real and checkable, rather than by an institution's word or an unverified mechanism. That ambition is not achieved in full here. Circulation, adoption, and liquidity — the conditions that took Bitcoin fifteen years of real use, not assertion, to establish — are not tested or claimed anywhere in this work. What is established is the requirement that neither fiat nor Bitcoin has actually delivered: a constraint that can be inspected, checked, and shown to hold under test, rather than taken on faith.")
doc.add_page_break()

# ===================== APPENDIX A =====================
add_chapter("Appendix A — Supplementary Empirical Checks")

add_heading("A.1 Differenced CEIR Boundary Check")
add_para("Chapter 4 reports that CEIR fails to predict returns once the analysis moves from levels to changes. This appendix reports that check in full. The differenced specification regresses 30-day forward returns on the change in log(CEIR) over the prior period, using the same HAC(30) standard errors as the level specification. The differenced coefficient is not statistically distinguishable from zero at conventional levels, in both the full sample and the pre-split subsample. This is consistent with the level result being driven by shared persistence in two slow-moving series — market capitalization and cumulative mining cost — rather than a genuine period-to-period causal relationship (Granger & Newbold, 1974). Augmented Dickey-Fuller tests do not reject a unit root in log(CEIR), log market capitalization, or log price, and Engle-Granger tests do not support a stable long-run equilibrium between market capitalization and either cumulative cost or cumulative electricity use (Dickey & Fuller, 1979; Engle & Granger, 1987).")

add_heading("A.2 Trading-Rule Negative Result")
add_para("As a practical-relevance check independent of the formal statistical tests, a simple trading rule was constructed directly from the CEIR signal: positioned according to whether log(CEIR) sits above or below its trailing median. Over the full sample period, this rule returns approximately +176.4%, compared with approximately +2,771.5% for simply holding Bitcoin throughout the same period. The corresponding Sharpe ratios are 0.72 for the CEIR rule and 1.13 for buy-and-hold (Lo, 2002). The ratio does not clear a basic practical-usefulness bar, independent of and in addition to its failure on all four formal tests in Chapter 4.")
add_table(
    ["Strategy", "Total return", "Sharpe ratio"],
    [["CEIR rule", "+176.4%", "0.72"],
     ["Buy-and-hold", "+2,771.5%", "1.13"]]
)
add_para("Table A.1. Trading-rule comparison over the full sample period.", italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ===================== APPENDIX B =====================
add_chapter("Appendix B — Implementation Evidence and Source Lineage")

add_heading("B.1 Sepolia Execution")
add_para("The proof-of-concept implementation described in Chapter 6 executed selected issuance and settlement rules on Ethereum's Sepolia public test network. A signed evidence bundle covering 2,606.7 kWh of accepted surplus production resulted in 130.1697 SPK minted, at a stated $0.05/kWh conversion rate used for this run. Source and attestation hashes are consumed once used, preventing the same evidence from supporting a second issuance. Invoice hashes prevent the same payment from being replayed. Redemption accounting converts spent tokens into owed-kWh claims that can later be marked fulfilled, in shortfall, or disputed.")

add_heading("B.2 Controlled Case Pack")
add_para("The three controlled cases reported in Chapter 6 — Taoyuan, Austin, and Phoenix — use model-based production data (PVWatts, with typical-meteorological-year context), with provider identity, coordinates, system size, annual output, data source, and a cryptographic hash preserved together for each case. The Taoyuan counterfactual holds the evidence file and its hash constant across both runs; only the declared assurance state changes, from a level that fails the pilot policy's minimum-provenance requirement to one that satisfies it.")

add_heading("B.3 Test Suite Results")
add_para("At the reviewed source revision, the deterministic decision core passed 60 of 60 tests, and the front-end passed 64 of 64. Additional checks at the same revision — production build, browser walkthrough, continuous integration, smart-contract-specific checks, security checks, and secrets checks — also passed. These results support internal consistency of the implementation at the reviewed revision. They do not substitute for independent operator validation or commercial testing.")

add_heading("B.4 Source-Lineage Audit")
add_para("The empirical analysis in Chapters 3 and 4 uses a local Bitcoin price export, an approximate supply curve, Cambridge Bitcoin Electricity Consumption Index exports, Cambridge Mining Map context, and the Alternative.me Fear and Greed Index. The upstream vendor and request metadata for the price export were not preserved, which limits exact replay of the original data pull. The intended geography-weighted electricity-price merge, described in Chapter 3, did not complete as designed; the stored series is near-constant at approximately $0.076/kWh, with a small number of manual overrides, rather than the intended monthly geography-weighted path. This limitation is disclosed directly in Chapters 3, 4, and 7, and does not affect the negative-control comparisons, which do not depend on the electricity-price series being accurate — only on the cumulative-cost and cumulative-electricity-use series being correctly constructed from the underlying CBECI data.")

add_heading("B.5 Reproduction Reference")
add_para("The frozen case-workbench revision referenced throughout Chapter 6 is recorded as: eb8714a6544b3480226283a69d41b3946df63451. Data and code notes for the empirical panel in Chapters 3 and 4 are maintained alongside the regression scripts referenced in the project's source repository.")
doc.add_page_break()

# ===================== REFERENCES =====================
add_chapter("References")
refs = [
"Bank for International Settlements. (2023). Blueprint for the future monetary system: Improving the old, enabling the new. In Annual Economic Report 2023. Bank for International Settlements.",
"Barro, R. J., & Gordon, D. B. (1983). Rules, discretion and reputation in a model of monetary policy. Journal of Monetary Economics, 12(1), 101-121.",
"Black, F., & Scholes, M. (1973). The pricing of options and corporate liabilities. Journal of Political Economy, 81(3), 637-654.",
"Bordo, M. D. (1993). The gold standard, Bretton Woods and other monetary regimes: A historical appraisal. Federal Reserve Bank of St. Louis Review, 75(2), 123-191.",
"Cambridge Centre for Alternative Finance. (n.d.-a). Cambridge Bitcoin Electricity Consumption Index: Methodology. Cambridge Judge Business School.",
"Cambridge Centre for Alternative Finance. (n.d.-b). Bitcoin Mining Map. Cambridge Judge Business School.",
"Cambridge Centre for Alternative Finance. (n.d.-c). Cambridge Bitcoin Electricity Consumption Index. Cambridge Judge Business School.",
"Chainlink. (2025). Decentralized oracle networks: Technical overview. Chainlink Labs.",
"Chow, G. C. (1960). Tests of equality between sets of coefficients in two linear regressions. Econometrica, 28(3), 591-605.",
"Cong, L. W., & He, Z. (2019). Blockchain disruption and smart contracts. The Review of Financial Studies, 32(5), 1754-1797.",
"Cox, J. C., Ross, S. A., & Rubinstein, M. (1979). Option pricing: A simplified approach. Journal of Financial Economics, 7(3), 229-263.",
"de Vries, A. (2018). Bitcoin's growing energy problem. Joule, 2(5), 801-805.",
"Dickey, D. A., & Fuller, W. A. (1979). Distribution of the estimators for autoregressive time series with a unit root. Journal of the American Statistical Association, 74(366a), 427-431.",
"Dobos, A. P. (2014). PVWatts Version 5 Manual (NREL/TP-6A20-62641). National Renewable Energy Laboratory.",
"Eichengreen, B. (1992). Golden fetters: The gold standard and the Great Depression, 1919-1939. Oxford University Press.",
"Engle, R. F., & Granger, C. W. J. (1987). Co-integration and error correction: Representation, estimation, and testing. Econometrica, 55(2), 251-276.",
"Eskandari, S., Salehi, M., Gu, W. C., & Clark, J. (2021). SoK: Oracles from the ground truth to market manipulation. arXiv:2106.00667.",
"Federal Reserve Bank of St. Louis. (2010). Central bank credibility and policy communication. Federal Reserve Bank of St. Louis Review.",
"Granger, C. W. J., & Newbold, P. (1974). Spurious regressions in econometrics. Journal of Econometrics, 2(2), 111-120.",
"Hayes, A. S. (2019). Bitcoin price and its marginal cost of production: Support for a fundamental value. Applied Economics Letters, 26(7), 554-560.",
"Hodrick, R. J. (1992). Dividend yields and expected stock returns: Alternative procedures for inference and measurement. The Review of Financial Studies, 5(3), 357-386.",
"International Energy Agency. (2023). Scaling up private finance for clean energy in emerging and developing economies.",
"Joskow, P. L. (2011). Comparing the costs of intermittent and dispatchable electricity generating technologies. American Economic Review, 101(3), 238-241.",
"Kiayias, A., & Lazos, P. (2022). SoK: Blockchain governance. arXiv:2201.07188.",
"Kristoufek, L. (2015). What are the main drivers of the Bitcoin price? Evidence from wavelet coherence analysis. PLOS ONE, 10(4).",
"Kronmal, R. A. (1993). Spurious correlation and the fallacy of the ratio standard revisited. Journal of the Royal Statistical Society: Series A, 156(3), 379-392.",
"Kydland, F. E., & Prescott, E. C. (1977). Rules rather than discretion: The inconsistency of optimal plans. Journal of Political Economy, 85(3), 473-491.",
"Lazard. (2025). Levelized cost of energy analysis. Lazard.",
"Liu, Y., & Tsyvinski, A. (2021). Risks and returns of cryptocurrency. The Review of Financial Studies, 34(6), 2689-2727.",
"Lo, A. W. (2002). The statistics of Sharpe ratios. Financial Analysts Journal, 58(4), 36-52.",
"Lucia, J. J., & Schwartz, E. S. (2002). Electricity prices and power derivatives: Evidence from the Nordic Power Exchange. Review of Derivatives Research, 5(1), 5-50.",
"Nakamoto, S. (2008). Bitcoin: A peer-to-peer electronic cash system.",
"NASA POWER. (n.d.). Daily API. NASA Langley Research Center.",
"Newey, W. K., & West, K. D. (1987). A simple, positive semi-definite, heteroskedasticity and autocorrelation consistent covariance matrix. Econometrica, 55(3), 703-708.",
"Sengupta, M., Habte, A., Wilbert, S., Gueymard, C. A., Remund, J., Lorenz, E., van Sark, W., & Jensen, A. R. (2024). Best practices handbook for the collection and use of solar resource data for solar energy applications (4th ed.; NREL/TP-5D00-88300). National Renewable Energy Laboratory.",
"Ueckerdt, F., Hirth, L., Luderer, G., & Edenhofer, O. (2013). System LCOE: What are the costs of variable renewables? Energy, 63, 61-75.",
"Zhang, F., Cecchetti, E., Croman, K., Juels, A., & Shi, E. (2016). Town Crier: An authenticated data feed for smart contracts. Proceedings of the 2016 ACM SIGSAC Conference on Computer and Communications Security, 270-282.",
]
for ref in refs:
    p = doc.add_paragraph(); r = p.add_run(ref)
    r.font.size=Pt(10.5); r.font.name='Times New Roman'
    p.paragraph_format.left_indent=Inches(0.5); p.paragraph_format.first_line_indent=Inches(-0.5)
    p.paragraph_format.space_after=Pt(10)

doc.save('/tmp/constrained_ledger_final.docx')
print("Saved.")
