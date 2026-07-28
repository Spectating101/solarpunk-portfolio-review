from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom'):
        if edge in kwargs:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(kwargs[edge]))
            el.set(qn('w:color'), '000000')
            tcBorders.append(el)
    tcPr.append(tcBorders)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(7); style.paragraph_format.line_spacing = 1.32
for s in doc.sections:
    s.left_margin = Inches(1); s.right_margin = Inches(1); s.top_margin = Inches(1); s.bottom_margin = Inches(1)

def add_heading(text, size=13.5):
    h = doc.add_paragraph(); run = h.add_run(text)
    run.font.name='Times New Roman'; run.font.size=Pt(size); run.font.bold=True
    h.paragraph_format.space_before=Pt(14); h.paragraph_format.space_after=Pt(6)
    return h

def add_sub(text, size=12):
    h = doc.add_paragraph(); run = h.add_run(text)
    run.font.name='Times New Roman'; run.font.size=Pt(size); run.font.bold=True
    h.paragraph_format.space_before=Pt(10); h.paragraph_format.space_after=Pt(4)
    return h

def add_para(text, bold=False, italic=False, size=11, align=None, space_after=7):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic
    p.paragraph_format.space_after=Pt(space_after)
    if align: p.alignment = align
    return p

def add_quote(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(11.5); r.font.italic=True
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text=''; p=hdr[i].paragraphs[0]; r=p.add_run(h)
        r.font.bold=True; r.font.size=Pt(9); r.font.name='Times New Roman'
        set_cell_border(hdr[i], top=8, bottom=4)
    for row in rows:
        cells = t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; p=cells[i].paragraphs[0]; r=p.add_run(str(val))
            r.font.size=Pt(9); r.font.name='Times New Roman'
    for c in t.rows[-1].cells: set_cell_border(c, bottom=8)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return t

def add_figure(path, caption, width=5.6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.font.size=Pt(9); r.font.italic=True; r.font.name='Times New Roman'
    cap.paragraph_format.space_after=Pt(12)

FIG = "/tmp/cl_figures"

# ===================== TITLE =====================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("The Constrained Ledger:\nWhen Does Energy Actually Discipline Digital Money?")
r.font.size=Pt(16); r.font.bold=True; r.font.name='Times New Roman'
tp.paragraph_format.space_after=Pt(16)
for line,sz,bold in [("Christopher Ongko (王新福)",12.5,True),
    ("Department of Finance, College of Management, Yuan Ze University, Taiwan",10.5,False),
    ("s1133958@mail.yzu.edu.tw · ORCID: 0009-0007-9339-9098",10.5,False),
    ("Working paper draft · July 2026",10.5,False)]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=p.add_run(line); rr.font.size=Pt(sz); rr.font.bold=bold; rr.font.name='Times New Roman'
    p.paragraph_format.space_after=Pt(3)
doc.add_paragraph().paragraph_format.space_after=Pt(6)

# ===================== ABSTRACT =====================
add_heading("Abstract", size=12.5)
for para in [
"Bitcoin was engineered on a specific claim: that a currency's value could be disciplined by real, verifiable cost rather than institutional trust. This paper tests whether that claim was ever actually delivered, using China's June 2021 mining ban as the clearest available natural experiment. A widely cited ratio of market value to cumulative mining cost (Hayes, 2019) is tested against four falsifiable criteria stated in advance: it must survive substitution with unrelated cumulative variables, survive a change in starting assumptions, hold in changes as well as levels, and be confirmed by a correctly specified structural-break test. The ratio fails all four. It is nearly indistinguishable from a ratio built on cumulative electricity use (correlation of .999989) and is similarly matched by a ratio built on cumulative elapsed time, depends materially on an arbitrary pre-sample accounting choice, disappears in differenced form, and receives no support from the preferred joint robust break test.",
"This negative result motivates the paper's constructive contribution. I specify five conditions as jointly necessary for an energy-linked financial claim to be genuinely constrained: reliable evidence, rule-bound issuance, explicit pricing of uncertainty, defined settlement, and limited governance. A transparent pricing framework, calibrated to five years of NASA satellite irradiance data, converts renewable-energy uncertainty into inspectable option prices and margin requirements. I then test a working rule system implementing all five conditions against the same four-part standard that eliminated the passive shortcut — evaluated across differentiated cases, under a deliberate counterfactual, under changing settlement conditions, and through execution on a public blockchain test network — and it satisfies all four criteria.",
"I do not claim to have built a functioning currency. Circulation, adoption, and liquidity are not tested or claimed. What I demonstrate is the piece that neither fiat money nor Bitcoin has actually delivered: a financial constraint tied to real evidence that can be inspected and shown to hold under test, rather than taken on faith.",
]:
    add_para(para)
add_para("Keywords: energy-linked finance, digital money, monetary credibility, Bitcoin energy cost, renewable-energy risk, smart contracts, reproducible decision systems", italic=True, size=9.5)
add_para("JEL Codes: E42, G13, Q42, Q47", italic=True, size=9.5)

# ===================== 1. INTRODUCTION =====================
add_heading("1. Introduction")
add_para("Bitcoin was engineered, not only launched, on a specific claim: that a currency's value could be disciplined by real, verifiable cost rather than institutional trust. In June 2021, that claim faced its clearest available test. China's State Council banned cryptocurrency mining nationwide, forcing roughly seventy percent of Bitcoin's global computing power to relocate within months. If cumulative mining cost genuinely anchors Bitcoin's price the way its engineering was intended to guarantee, this shock should have produced a detectable change in that relationship.")
add_para("A specific ratio — market value divided by cumulative mining expenditure — has circulated in crypto analysis and is defended in published research (Hayes, 2019) as evidence that this claim held. Tested directly against daily data from 2019 through 2025, it does not survive scrutiny. A formal test for whether it changed after the ban finds no reliable break, and replacing cumulative electricity cost with an unrelated variable — cumulative elapsed time — produces a nearly identical result. The ratio was never capturing information specific to electricity price, and this finding is the starting point for the analysis that follows.")
add_para("Bitcoin's own founding document explicitly stated its ambition: its title called it “a peer-to-peer electronic cash system” — a currency, declared as a goal, years before any real adoption existed to support that claim. I take up that same ambition here: not a narrower financial instrument for its own sake, but the objective Bitcoin itself declared and, as I demonstrate, did not deliver — a currency system whose value is genuinely disciplined by something real and checkable, rather than by an institution's word or an unverified mechanism. The ambition is not achieved in full here; no claim is made about circulation, adoption, or liquidity, which Bitcoin itself established only empirically, over fifteen years of real use. What I demonstrate is the specific requirement needed to make that ambition achievable at all, tested directly for the first time.")
add_para("An energy reference does not, by itself, constrain anything. A fixed token supply does not reveal whether the underlying evidence is reliable, whether the permitted quantity is defensible, or whether settlement obligations can be met. The question I address is:")
add_quote("Under what conditions can verified energy evidence impose a credible, checkable constraint on a digital financial claim?")
add_para("I answer it in three connected steps: testing the strongest available shortcut against a real natural experiment; translating renewable-energy uncertainty into concrete, priced financial quantities; and specifying and implementing the conditions under which those quantities could actually govern a real claim, tested against controlled cases with differentiated, traceable outcomes.")
add_sub("1.1 Contributions", size=11.5)
add_para("First, I identify the limits of Bitcoin's cumulative mining cost as a value anchor: the relationship a widely cited ratio (Hayes, 2019) claims to show is not specific to electricity price, does not survive a proper structural-break test, and disappears under differencing.")
add_para("Second, I build a transparent pricing framework for renewable-energy uncertainty, calibrated to five years of satellite irradiance data, converting a physically variable resource into inspectable option prices and margin requirements.")
add_para("Third, I specify five jointly necessary conditions for an energy-linked financial claim to be credible — reliable evidence, rule-bound issuance, explicit pricing of uncertainty, defined settlement, and limited governance — and implement and test a system built on them, including execution on a public blockchain and controlled test cases with traceable, differentiated outcomes.")
add_sub("1.2 Scope", size=11.5)
add_para("The claims here are deliberately bounded. I do not argue that energy should replace fiat money, do not treat Bitcoin as energy-backed, and do not present the implemented system as legally or commercially ready. The Bitcoin analysis identifies the limits of one mining-cost ratio, not a general law of digital-asset valuation. The renewable-energy model is a cold-start framework for exposing pricing and risk assumptions, not a traded price or legal settlement mechanism. The implementation is proof-of-concept research on a public testnet, not evidence of operator validation or deployment readiness.")

# ===================== 2. THEORETICAL FRAMEWORK =====================
add_heading("2. Why an Energy Reference Is Not Automatically a Constraint")
add_para("The existing literature does not lack pieces of this puzzle. It has real explanations for scarcity, credibility, production risk, and automated enforcement. What it lacks is an account of how those pieces interact before a physical reference can actually limit a digital financial claim.")
add_sub("2.1 Monetary Credibility, Rules, and Discretion", size=11.5)
add_para("Kydland and Prescott (1977) demonstrate that a rule announced today can be abandoned once people have already acted on it, so credibility depends not only on the current rule but on how costly and visible it would be to change it later. Barro and Gordon (1983) identify the related expectations problem: promises are discounted whenever people expect future discretion. In digital finance, a claim of limited supply or energy backing remains weak for the same reason if administrators can mint outside the rule, replace the evidence source, or avoid redemption when inconvenient. The design principle this yields is not a specific asset choice — it is that the limits must bind the issuer as much as the user.")
add_sub("2.2 Gold, Fiat, and Bitcoin", size=11.5)
add_para("Gold-backed money combined physical scarcity with a working redemption promise; its most visible failure, at Bretton Woods, was not gold disappearing — it was the arithmetic connecting outstanding claims to settlement capacity breaking down (Bordo, 1993; Eichengreen, 1992). Fiat money replaced that physical promise with institutional authority and a track record, trading flexibility for dependence on political will and repeated performance (Federal Reserve Bank of St. Louis, 2010). Energy differs from both in ways that matter throughout the remainder of this paper: it is time- and location-dependent, often impossible to store economically, and inseparable from the infrastructure needed to deliver it. An energy-linked system cannot import gold's convertibility or fiat's institutional depth by analogy — it must define its own evidence, pricing, settlement, and governance.")
add_para("Bitcoin's issuance schedule is fixed by protocol, and proof-of-work makes rewriting transaction history costly (Nakamoto, 2008). It is neither redeemable like gold nor administered like fiat. Proof-of-work does give digital scarcity a real expenditure base: miners consume electricity and hardware to compete for rewards, making Bitcoin genuinely costly to produce and secure. That cost is not redeemable backing, however — holders receive no right to the electricity actually consumed (de Vries, 2018). This makes Bitcoin a useful boundary case. Any distinctive mining-expenditure effect on price must be separated from ordinary crypto return dynamics, price persistence, and how the ratio itself is constructed (Granger & Newbold, 1974; Kristoufek, 2015; Kronmal, 1993; Liu & Tsyvinski, 2021) — exactly what Section 3 tests against the production-cost claim introduced above (Hayes, 2019). The available cost data has real limits, too: the Cambridge Bitcoin Electricity Consumption Index is model-based, not direct metering (Cambridge Centre for Alternative Finance, n.d.-a), and any cumulative-cost series inherits that uncertainty. This is why the empirical test in Section 3 treats the resulting ratio as falsifiable rather than as a measure of intrinsic value.")
add_sub("2.3 Renewable-Energy Finance and Pricing", size=11.5)
add_para("Solar and wind assets generate electricity with direct economic use, but the financial value of that output depends on weather, location, timing, grid access, and settlement rules (Joskow, 2011; Ueckerdt et al., 2013). Cost competitiveness does not eliminate those contingencies. Public data illustrates a real evidence hierarchy: NASA POWER and PVWatts are useful for screening and scenario analysis, but neither can prove that a specific site generated, exported, or settled a given quantity of electricity (Dobos, 2014; Sengupta et al., 2024). Resource potential, expected output, observed generation, and settled delivery are genuinely different things, and Section 4 treats them accordingly. Pricing theory forces that uncertainty into explicit, checkable assumptions rather than leaving it implicit. Black and Scholes (1973) and Cox, Ross, and Rubinstein (1979) are used in Section 4 for this reason — not because electricity behaves like an ordinary stock, but because both require the analyst to declare the underlying proxy, volatility, horizon, and payoff explicitly. Standard option assumptions fit electricity imperfectly, since power prices can spike, mean-revert, and turn negative (Lucia & Schwartz, 2002). Option-style reasoning is used here as a transparent benchmark, not a complete market model.")
add_sub("2.4 Smart Contracts, Oracles, and Governance", size=11.5)
add_para("Smart contracts can enforce issuance and settlement rules, but execution quality is not economic quality — automation can apply a flawed rule or accept inaccurate data as faithfully as sound ones (Cong & He, 2019). The part that matters directly for what follows is the oracle problem: a blockchain cannot itself observe whether electricity was generated or delivered, so external data feeds must supply that fact. This introduces its own reliability and manipulation risk (Chainlink, 2025; Eskandari et al., 2021; Zhang et al., 2016), a limitation Section 6 addresses directly. Governance introduces a parallel risk: an administrator can defeat every other safeguard without falsifying the data, simply by minting outside accepted evidence (Kiayias & Lazos, 2022). This framework is not a stablecoin design. It addresses an earlier question than reserve policy or peg mechanics — whether verified evidence can restrict claim creation at all (Bank for International Settlements, 2023).")
add_para("Credibility, across this literature, emerges from a chain rather than from any single asset or technology. Bitcoin research studies mining expenditure without connecting it to a direct energy claim; renewable-energy finance studies production without specifying issuance rules; pricing research studies uncertain payoffs without specifying how the result should limit claim creation; smart-contract research studies enforcement without determining whether the underlying claim is economically defensible. I answer that gap with five constraints — reliable evidence, rule-bound issuance, explicit pricing of uncertainty, defined settlement, and limited governance — which the remaining sections test, price, and implement directly.")

# ===================== 3. TESTING BITCOIN =====================
add_heading("3. Testing Bitcoin as the Clearest Case")
add_para("In this section, I test the shortcut that motivates the paper: whether real mining expenditure already anchors Bitcoin's value on its own. Bitcoin is the relevant case because proof-of-work converts electricity and computation into one measurable, cumulative cost. For that shortcut to count as genuine evidence of an energy-price constraint, rather than a suggestive number, it must clear four tests, stated here before any result is reported: (1) it must survive substitution with unrelated cumulative quantities; (2) it must survive a change in starting assumptions; (3) it must hold in changes, not only levels; (4) it must be confirmed by the correctly specified test where a simpler one disagrees. This is the same standard the constructed system in Section 5 is later held to.")
add_para("The test is organized around the Cumulative Energy Investment Ratio:")
add_quote("CEIRₜ = Market Capitalisationₜ / Cumulative Energy Costₜ")
add_para("A high value indicates Bitcoin's market value is large relative to estimated cumulative mining cost; the candidate sign is negative. Because the denominator accumulates slowly and persistently, this interpretation is vulnerable to spurious regression (Granger & Newbold, 1974; Kronmal, 1993). The panel runs from January 2019 through May 2025, with 2,280 observations once forward returns and controls are complete. Bitcoin price and market capitalization are derived from a local daily price export and an approximate supply curve; mining electricity use comes from the Cambridge Bitcoin Electricity Consumption Index, a model-based estimate (Cambridge Centre for Alternative Finance, n.d.-a). The intended geography-weighted electricity-price series did not merge as designed, leaving a near-constant price path in its place. The candidate split date, June 20, 2021, is a researcher-defined proxy for the China mining crackdown. Prediction uses HAC(30) standard errors for overlapping return windows (Hodrick, 1992; Newey & West, 1987). The regime-change test uses a joint robust break test alongside the classical Chow statistic (Chow, 1960). The mechanism test removes the starting cost assumption, substitutes cumulative electricity use and elapsed days for cumulative cost, and checks changes as well as levels (Dickey & Fuller, 1979; Engle & Granger, 1987).")
add_table(
    ["Series", "Source", "Role"],
    [["Bitcoin price and derived market capitalization", "Local daily price export; approximate BTC supply curve", "Return outcome; CEIR numerator"],
     ["Mining electricity", "Cambridge CBECI best-guess annualized TWh", "Energy-cost base"],
     ["Legacy electricity-price input", "Intended geography-weighted merge; failed, near-constant instead", "Legacy cost denominator"],
     ["Mining geography", "Cambridge Mining Map, monthly shares", "Context for the candidate split"]]
)
add_para("Table A. Data sources for the empirical panel.", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table(
    ["Variable", "Mean (full)", "Mean (pre-split)", "Mean (post-split)"],
    [["Bitcoin price (USD)", "34,881", "16,666", "46,161"],
     ["CEIR", "29.65", "30.40", "29.19"],
     ["30-day forward return", "6.4%", "10.7%", "3.8%"]]
)
add_para("Table B. Descriptive statistics, regression sample (N = 2,280). Bitcoin's mean price roughly tripled after the split while mean CEIR stayed in a broadly similar range — the two series do not move in lockstep, which is part of why formal inference, rather than the raw picture, must carry the argument below.", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
add_sub("3.1 Results")
add_para("The level regression reproduces the attractive result: before the candidate split, a higher CEIR predicts weaker 30-day returns, with a coefficient of -0.262 (HAC p < .001). Table 1 reports what happens under each of the four tests.")
add_table(
    ["Test", "Result", "Reading"],
    [["1. Substitution", "Correlation of .999989 with a cumulative-electricity-use version; similarly strong with a day-count version", "Fails — not specific to price"],
     ["2. Starting assumptions", "Removing ~$3.30B of pre-sample cost weakens the coefficient from -0.262 to -0.104 (p rises from <.001 to .099)", "Fails — depends on an arbitrary choice"],
     ["3. Changes vs. levels", "No reliable predictive relationship in differenced form", "Fails — consistent with persistence"],
     ["4. Correct break test", "Classical Chow suggests a break; joint robust test does not (p = .133)", "Fails to confirm a break"]]
)
add_para("Table 1. CEIR against the four-part standard. Post-split coefficient: β = -0.0708, p = .133 (N = 1,408); pre-split: β = -0.2623, p = .00052 (N = 872).", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Two further checks reach the same conclusion from different angles. Converting the level specification into short-run changes removes the apparent predictive effect entirely: differenced log(CEIR) does not reliably predict later returns, consistent with the level result being driven by shared persistence in two slow-moving series rather than a genuine period-to-period relationship (Granger & Newbold, 1974). Augmented Dickey-Fuller tests do not reject a unit root in log(CEIR), log market capitalization, or log price, and Engle-Granger tests do not support a stable long-run equilibrium between market capitalization and either cumulative cost or cumulative electricity use (Dickey & Fuller, 1979; Engle & Granger, 1987). As a practical-relevance check independent of the formal tests, a simple trading rule built directly on the CEIR signal returns approximately +176.4% over the full sample, against approximately +2,771.5% for simply holding Bitcoin; the corresponding Sharpe ratios are 0.72 and 1.13 (Lo, 2002). The ratio does not clear a basic practical-usefulness bar, on top of failing all four formal tests.")
add_para("CEIR fails all four tests. Mining expenditure is real; this ratio does not isolate its financial effect. This is not a disappointing result: even had CEIR passed every test, a single ratio cannot specify an evidence standard, an issuance rule, a settlement obligation, or a governance limit — at most it would have been indirect evidence that discipline was present, never the discipline itself. In what follows, I build the two components a real constraint requires: a way to price genuine uncertainty (Section 4), and an implemented rule system tested against the same four-part standard (Section 5).")

# ===================== 4. PRICING =====================
add_heading("4. Pricing Renewable-Energy Uncertainty")
add_para("Renewable energy is genuinely productive, but its financial value depends on weather, location, and time of day — a forward-looking, ongoing uncertainty unlike Bitcoin's fixed, backward-looking cumulative cost. No liquid options market exists for this kind of contract to check a price against, so I build a defensible number from first principles, using five years of daily satellite-measured solar irradiance data (NASA POWER) for a representative Taiwan location. Raw daily irradiance is smoothed with a four-day rolling average before computing daily log returns, with the top and bottom one percent trimmed as measurement outliers, producing an annualized volatility estimate of approximately 189 percent — high, but reflecting direct physical output variability rather than financial market volatility.")
add_para("That volatility feeds a standard option-pricing setup — Geometric Brownian Motion, priced two independent ways: a binomial tree (Cox, Ross, & Rubinstein, 1979) and a Monte Carlo simulation, chosen because both force every assumption to be stated explicitly.")
add_table(
    ["Input", "Value"],
    [["Spot / strike price", "$0.0516/kWh"],
     ["Annualized volatility", "~189%"],
     ["Binomial tree price (N=400)", "$0.0356/kWh"],
     ["Monte Carlo price", "$0.0361/kWh"],
     ["Agreement between methods", "Within 2%"]]
)
add_para("Table 2. Taiwan base-case pricing.", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("The two methods agreeing within two percent is a consistency check, not proof of a market-correct price, since none yet exists to validate against. The same volatility estimate produces a margin requirement of approximately $0.63/kWh and a data-tolerance bound of approximately 21.7 percent. Only Taiwan's volatility is calibrated from fully disclosed real data; the pricing tool built here is an input to Section 5's system, not a market-validated instrument in its own right.")

# ===================== 5. THE SYSTEM =====================
add_heading("5. The Built, Tested System")
add_para("With the five conditions specified and a pricing method available, I implement a working rule system in this section rather than leaving the conditions as description, separating four decisions: whether the evidence is trustworthy, whether policy allows the claim, how large it may be, and whether it can be settled. Submitted evidence is checked against a declared policy and hashed to prevent reuse; admission returns a plain PASS or BLOCK; quantity is capped by whichever applicable limit is lowest; settlement is evaluated later against whatever capacity is actually available. A receipt preserves the inputs, the rule set, the binding limit, and the outcome.")
add_figure(f"{FIG}/fig2_pipeline.png", "Figure 1. The four-gate decision pipeline, applied to three real test cases with three different outcomes.")
add_para("This section applies the tests specified in Section 3 to the constructed system.")
add_table(
    ["Section 3's test", "Applied here as", "Result"],
    [["1. Substitution", "Three different cases — Taoyuan, Austin, Phoenix", "Each produces a different, correctly traced outcome"],
     ["2. Starting assumptions", "A counterfactual: only the assurance level changes, evidence untouched", "Outcome changes from blocked to admitted, tracing to the rule that changed"],
     ["3. Changing conditions", "Settlement tested against three capacity scenarios", "Correctly reports full settlement, partial settlement, and shortfall"],
     ["4. Correct test", "Real execution on a public blockchain test network", "Produces an independently checkable transaction"]]
)
add_para("Table 3. The four-part standard, applied to the constructed system.", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Taoyuan is submitted with evidence that falls short of the required standard and is blocked outright, with no quantity calculated. In the counterfactual, the evidence file is untouched and only the declared assurance level is raised; admission then passes, and the resulting quantity — 126 units — traces directly to the provenance rule that changed. Austin is admitted at 283.09811 units because the resource data itself is the lowest applicable ceiling; Phoenix is admitted at up to 320 units because a different, evidence-backed capacity constraint binds instead. Run against three settlement-capacity scenarios, the system correctly reports full settlement, partial settlement (50.4 units covered, 75.6 in shortfall, at 40 percent capacity), and outright shortfall. One execution ran on Ethereum's Sepolia test network, minting SPK against 2,606.7 kWh of accepted surplus, producing a real, independently checkable transaction. At the reviewed revision, the deterministic core passed 60 of 60 tests and the front end passed 64 of 64.")
add_para("Where CEIR failed all four of Section 3's tests, this system satisfies the corresponding four criteria. This establishes that evidence, issuance, pricing, settlement, and governance can be represented as separate, inspectable decisions that change outcomes, not only how they are explained. It does not establish that the underlying data sources are production-grade — satellite estimates are not metered output, and oracles carry their own manipulation risk — nor legal enforceability, real market demand, or production-grade security. This is proof-of-concept evidence under controlled conditions.")

# ===================== 6. DISCUSSION =====================
add_heading("6. Discussion and Conclusion")
add_sub("6.1 Answer to the Research Question", size=11.5)
add_para("The question I opened with was under what conditions verified energy evidence can impose a credible, checkable constraint on a digital financial claim. The evidence developed across Sections 3 through 5 supports a conditional answer. Passive expenditure alone does not constitute such a constraint — Bitcoin's cumulative mining cost fails every test that would distinguish a genuine energy-price effect from a persistent, coincidental valuation pattern. A constraint becomes credible only when five conditions are held jointly: reliable evidence, rule-bound issuance, explicit pricing of uncertainty, defined settlement, and limited governance. Sections 4 and 5 demonstrate that these conditions can be implemented, and that the resulting system behaves correctly under the same discipline that eliminated the passive shortcut.")
add_sub("6.2 Limitations", size=11.5)
add_para("Several limitations qualify these findings directly. The empirical evidence in Section 3 is limited to one proof-of-work asset, using a market-capitalization series derived from a local price export and an approximate supply curve rather than a fully documented observed series. Mining electricity use is itself a model-based estimate, not direct metering, and the intended geography-weighted electricity-price series did not merge as designed. These limitations do not prevent reproduction of the original result; they limit how strongly that result, or its failure, can be generalized to a differently constructed cost measure.")
add_para("The pricing framework in Section 4 remains a cold-start scenario model. Real power markets exhibit jumps, negative prices, mean reversion, and seasonal patterns beyond what a Geometric Brownian Motion benchmark captures. Only the Taiwan case has a fully disclosed, preserved volatility calibration.")
add_para("The implementation in Section 5 is controlled proof-of-concept evidence. Execution on a public test network does not establish revenue-grade metering, legal enforceability, real market demand, or production-grade security. The satellite and model-based evidence used in the controlled cases is not equivalent to direct physical metering, and oracle risk — the possibility that external data feeds misreport or are manipulated — remains a real, unresolved boundary, consistent with Section 2.")
add_sub("6.3 What Would Strengthen or Weaken This Work", size=11.5)
add_para("The architecture developed here is falsifiable, and it is worth stating plainly what would count against it. Five things would weaken it: real operator evidence that could not be classified consistently under the stated policy; pricing outputs that could not be converted into defensible quantity limits under real market data; the same policy producing unpredictably different admission outcomes under equivalent evidence; settlement rights that remained ambiguous once tested against a real dispute; or governance that could waive the constraint whenever it became costly to honor. Conversely, the negative result on Bitcoin's mining cost would be weakened if a differently constructed, geography-accurate cost measure produced a result that survived the same four tests applied in Section 3. Two extensions would test these boundaries directly. A future extension would replace the model-based electricity-price series with a properly merged, geography-weighted series and repeat the identification tests under that correction. A pilot deployment with real operator evidence and real settlement obligations would test the implementation in Section 5 under conditions this proof-of-concept does not reach.")
add_sub("6.4 Conclusion", size=11.5)
add_para("China's 2021 mining ban was, in principle, the clearest available test of whether Bitcoin's cost of production disciplines its value. It does not. The relationship claimed by a widely cited production-cost ratio fails every test designed to distinguish a genuine energy-price effect from a coincidental one. What replaces that claim is not a theory but a system: five jointly necessary conditions, derived from how monetary discipline has actually functioned in gold-backed and fiat regimes, implemented as a working rule set, and tested against real cases with differentiated, traceable outcomes. The ambition I stated at the outset of this paper — a currency genuinely disciplined by something real and checkable, rather than by an institution's word or an unverified mechanism — is not achieved in full. Circulation, adoption, and liquidity are not tested or claimed anywhere in this work. What I establish is the requirement that neither fiat nor Bitcoin has actually delivered: a constraint that can be inspected, checked, and shown to hold under test, rather than taken on faith.")

# ===================== DATA STATEMENT =====================
add_heading("Data and Code Statement", size=12.5)
add_para("The empirical analysis uses a locally reconstructed Bitcoin price export, an approximate supply curve, Cambridge Bitcoin Electricity Consumption Index exports, Cambridge Mining Map context, and the Alternative.me Fear and Greed Index; source-lineage limitations, including the failed geography-weighted electricity-price merge, are disclosed directly in Section 3. The pricing analysis uses NASA POWER irradiance data and archived binomial and Monte Carlo scripts. The implementation is documented in the SolarPunk project artifacts, including Sepolia runtime records, controlled case-pack outcomes, and validation logs, frozen at revision eb8714a6544b3480226283a69d41b3946df63451. Full differencing and trading-rule results, and additional source-lineage detail, are available in the replication materials upon request.")

doc.add_page_break()
add_heading("References", size=13)
refs = [
"Bank for International Settlements. (2023). Blueprint for the future monetary system: Improving the old, enabling the new. In Annual Economic Report 2023.",
"Barro, R. J., & Gordon, D. B. (1983). Rules, discretion and reputation in a model of monetary policy. Journal of Monetary Economics, 12(1), 101-121.",
"Black, F., & Scholes, M. (1973). The pricing of options and corporate liabilities. Journal of Political Economy, 81(3), 637-654.",
"Bordo, M. D. (1993). The gold standard, Bretton Woods and other monetary regimes: A historical appraisal. Federal Reserve Bank of St. Louis Review, 75(2), 123-191.",
"Cambridge Centre for Alternative Finance. (n.d.-a). Cambridge Bitcoin Electricity Consumption Index: Methodology. Cambridge Judge Business School.",
"Chainlink. (2025). Decentralized oracle networks: Technical overview. Chainlink Labs.",
"Chow, G. C. (1960). Tests of equality between sets of coefficients in two linear regressions. Econometrica, 28(3), 591-605.",
"Cong, L. W., & He, Z. (2019). Blockchain disruption and smart contracts. The Review of Financial Studies, 32(5), 1754-1797.",
"Cox, J. C., Ross, S. A., & Rubinstein, M. (1979). Option pricing: A simplified approach. Journal of Financial Economics, 7(3), 229-263.",
"de Vries, A. (2018). Bitcoin's growing energy problem. Joule, 2(5), 801-805.",
"Dickey, D. A., & Fuller, W. A. (1979). Distribution of the estimators for autoregressive time series with a unit root. Journal of the American Statistical Association, 74(366a), 427-431.",
"Dobos, A. P. (2014). PVWatts Version 5 Manual (NREL/TP-6A20-62641). National Renewable Energy Laboratory.",
"Eichengreen, B. (1992). Golden fetters: The gold standard and the Great Depression, 1919-1939. Oxford University Press.",
"Engle, R. F., & Granger, C. W. J. (1987). Co-integration and error correction: Representation, estimation, and testing. Econometrica, 55(2), 251-276.",
"Eskandari, S., Salehi, M., Gu, W. C., & Clark, J. (2021). SoK: Oracles from the ground truth to market manipulation. arXiv:2106.00667.",
"Federal Reserve Bank of St. Louis. (2010). Central bank credibility and policy communication. Federal Reserve Bank of St. Louis Review.",
"Granger, C. W. J., & Newbold, P. (1974). Spurious regressions in econometrics. Journal of Econometrics, 2(2), 111-120.",
"Hayes, A. S. (2019). Bitcoin price and its marginal cost of production: Support for a fundamental value. Applied Economics Letters, 26(7), 554-560.",
"Hodrick, R. J. (1992). Dividend yields and expected stock returns: Alternative procedures for inference and measurement. The Review of Financial Studies, 5(3), 357-386.",
"Joskow, P. L. (2011). Comparing the costs of intermittent and dispatchable electricity generating technologies. American Economic Review, 101(3), 238-241.",
"Kiayias, A., & Lazos, P. (2022). SoK: Blockchain governance. arXiv:2201.07188.",
"Kristoufek, L. (2015). What are the main drivers of the Bitcoin price? Evidence from wavelet coherence analysis. PLOS ONE, 10(4).",
"Kronmal, R. A. (1993). Spurious correlation and the fallacy of the ratio standard revisited. Journal of the Royal Statistical Society: Series A, 156(3), 379-392.",
"Kydland, F. E., & Prescott, E. C. (1977). Rules rather than discretion: The inconsistency of optimal plans. Journal of Political Economy, 85(3), 473-491.",
"Liu, Y., & Tsyvinski, A. (2021). Risks and returns of cryptocurrency. The Review of Financial Studies, 34(6), 2689-2727.",
"Lo, A. W. (2002). The statistics of Sharpe ratios. Financial Analysts Journal, 58(4), 36-52.",
"Nakamoto, S. (2008). Bitcoin: A peer-to-peer electronic cash system.",
"Newey, W. K., & West, K. D. (1987). A simple, positive semi-definite, heteroskedasticity and autocorrelation consistent covariance matrix. Econometrica, 55(3), 703-708.",
"Sengupta, M., Habte, A., Wilbert, S., Gueymard, C. A., Remund, J., Lorenz, E., van Sark, W., & Jensen, A. R. (2024). Best practices handbook for the collection and use of solar resource data for solar energy applications (4th ed.; NREL/TP-5D00-88300). National Renewable Energy Laboratory.",
"Ueckerdt, F., Hirth, L., Luderer, G., & Edenhofer, O. (2013). System LCOE: What are the costs of variable renewables? Energy, 63, 61-75.",
"Zhang, F., Cecchetti, E., Croman, K., Juels, A., & Shi, E. (2016). Town Crier: An authenticated data feed for smart contracts. Proceedings of the 2016 ACM SIGSAC Conference on Computer and Communications Security, 270-282.",
]
for ref in refs:
    p = doc.add_paragraph(); r = p.add_run(ref)
    r.font.size=Pt(10); r.font.name='Times New Roman'
    p.paragraph_format.left_indent=Inches(0.5); p.paragraph_format.first_line_indent=Inches(-0.5)
    p.paragraph_format.space_after=Pt(9)

doc.save('/tmp/constrained_ledger_journal.docx')
print("Saved.")
